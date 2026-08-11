# -*- coding: utf-8 -*-
"""Génère le rapport de thèse GP AI au format Word (.docx), avec le formatage
exigé par le référentiel (Calibri 12, marges 2,5 cm, interligne 1,5)."""
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

DOSSIER = os.path.dirname(os.path.abspath(__file__))
RACINE = os.path.join(DOSSIER, "..")
SORTIE = os.path.join(DOSSIER, "ISMAIL_YACINE_THESE.docx")

VIOLET = RGBColor(0x5B, 0x2C, 0x83)
GRIS = RGBColor(0x55, 0x55, 0x55)

# ------------------------------------------------------------------ utilitaires
def style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for i, taille in [(1, 16), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Calibri"
        h.font.size = Pt(taille)
        h.font.bold = True
        h.font.color.rgb = VIOLET
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

def marges(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)

def numeros_pages(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for tag, txt in [("begin", None), (None, "PAGE"), ("end", None)]:
        if tag:
            el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), tag); run._r.append(el)
        else:
            el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve"); el.text = txt
            run._r.append(el)

def titre(doc, texte, niveau=1):
    doc.add_heading(texte, level=niveau)

def para(doc, texte, gras=False, italique=False, centre=False):
    p = doc.add_paragraph()
    r = p.add_run(texte); r.bold = gras; r.italic = italique
    if centre: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def puce(doc, texte):
    doc.add_paragraph(texte, style="List Bullet")

def apercu(doc, texte):
    """Encadré 'à compléter' pour les parties nécessitant les recherches du candidat."""
    p = doc.add_paragraph()
    r = p.add_run("[À COMPLÉTER] " + texte)
    r.italic = True; r.font.color.rgb = GRIS

def tableau(doc, entetes, lignes):
    t = doc.add_table(rows=1, cols=len(entetes))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, e in enumerate(entetes):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(e); run.bold = True; run.font.size = Pt(10)
    for ligne in lignes:
        cells = t.add_row().cells
        for i, val in enumerate(ligne):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(10)
    doc.add_paragraph()
    return t

def sommaire(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    run._r.addprevious(fld)
    msg = doc.add_paragraph()
    r = msg.add_run("(Sommaire automatique : clic droit → « Mettre à jour les champs » dans Word.)")
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRIS

# ------------------------------------------------------------------ document
doc = Document()
style_base(doc); marges(doc); numeros_pages(doc)
dico = json.load(open(os.path.join(RACINE, "_dict.json"), encoding="utf-8"))

# ---- Page de présentation ----
for _ in range(3):
    doc.add_paragraph()
para(doc, "GP AI", gras=True, centre=True).runs[0].font.size = Pt(36)
para(doc, "Système d'aide à la décision pour l'orientation diagnostique",
     gras=True, centre=True).runs[0].font.size = Pt(18)
para(doc, "des troubles génétiques chez l'enfant", gras=True, centre=True).runs[0].font.size = Pt(18)
doc.add_paragraph()
para(doc, "Thèse professionnelle", italique=True, centre=True)
para(doc, "Mastère Data et Intelligence Artificielle — RNCP 37137", centre=True)
para(doc, "Chef de projet Data et Intelligence Artificielle", centre=True)
for _ in range(4):
    doc.add_paragraph()
para(doc, "Rédacteur : Yacine Ismail", centre=True)
para(doc, "Date de début de projet : janvier 2026", centre=True)
para(doc, "Nexa Digital School", centre=True)
para(doc, "[Insérer ici le logo Nexa Digital School]", italique=True, centre=True).runs[0].font.color.rgb = GRIS
doc.add_page_break()

# ---- Sommaire ----
titre(doc, "Sommaire", 1)
sommaire(doc)
doc.add_page_break()

# ================================================================ 1. ENTREPRISE
titre(doc, "1. Descriptif de l'entreprise", 1)

titre(doc, "1.1 Storytelling — histoire de l'entreprise", 2)
para(doc, "GénoPed est une startup MedTech française fondée en 2019 à Lyon, spécialisée dans le "
     "diagnostic génétique pédiatrique. Elle est née de la rencontre entre le Dr. Claire Morin, "
     "généticienne pédiatrique au CHU de Lyon, et Thomas Bergier, ingénieur en bioinformatique. "
     "Après des années passées à observer des familles attendre plusieurs mois — parfois plusieurs "
     "années — avant d'obtenir un diagnostic pour leur enfant, ils ont décidé de créer une structure "
     "capable de centraliser, numériser et accélérer le parcours diagnostique en génétique pédiatrique.")
para(doc, "Soutenue par un financement BPI France en 2021 et labellisée « Startup Santé » par le "
     "ministère de la Santé en 2022, GénoPed s'est progressivement imposée comme un acteur de "
     "référence du diagnostic génétique assisté, en partenariat avec 12 établissements hospitaliers "
     "en France et en Belgique. L'entreprise compte aujourd'hui 47 employés et traite environ "
     "4 800 dossiers par an.")

titre(doc, "1.2 Valeurs et missions", 2)
puce(doc, "Rapidité diagnostique : réduire l'errance diagnostique des familles, estimée à 5 ans en "
      "moyenne en France pour les maladies rares.")
puce(doc, "Accessibilité : rendre le diagnostic génétique accessible aux établissements ne disposant "
      "pas de généticien sur place.")
puce(doc, "Éthique médicale : « l'IA oriente, le médecin décide », avec une transparence totale sur "
      "les prédictions et leurs limites.")
puce(doc, "Conformité RGPD : données de santé hébergées sur serveur HDS certifié, pseudonymisation "
      "systématique.")

titre(doc, "1.3 Activité principale", 2)
para(doc, "GénoPed accompagne les établissements de santé dans l'orientation et la confirmation des "
     "diagnostics de troubles génétiques chez l'enfant. Son cœur de métier consiste à recueillir les "
     "dossiers cliniques, à les structurer dans un dossier patient numérique, puis à assister le "
     "généticien dans l'interprétation, jusqu'à la prescription des examens de confirmation.")

titre(doc, "1.4 Environnement économique et sociétal", 2)
apercu(doc, "Situer le secteur de la e-santé et du diagnostic des maladies rares en France : "
       "poids économique, dynamique du marché de la MedTech, enjeux de santé publique liés aux "
       "3 millions de personnes atteintes de maladies rares en France (source : Plan National "
       "Maladies Rares, à citer). Ajouter des chiffres récents (< 5 ans) avec sources en note de bas "
       "de page (Ministère de la Santé, Orphanet, INSEE, France Assos Santé).")

titre(doc, "1.5 Environnement technologique", 2)
para(doc, "L'infrastructure de GénoPed repose sur une base de données relationnelle PostgreSQL "
     "hébergée sur une infrastructure certifiée HDS (Hébergeur de Données de Santé). Les traitements "
     "de données et l'intelligence artificielle sont développés en Python (Pandas, scikit-learn), et "
     "les applications métier sont exposées via des services web (Flask, API REST).")

titre(doc, "1.6 Environnement de données", 2)
para(doc, "Les dossiers patients sont centralisés dans la base PostgreSQL. C'est de cette base "
     "qu'est extrait, après pseudonymisation, le jeu de données ayant servi à l'apprentissage du "
     "modèle : 18 943 dossiers pédiatriques décrits par 30 variables cliniques (symptômes, "
     "antécédents familiaux, mesures biologiques) et une variable cible (la sous-classe "
     "diagnostique, parmi 9 troubles génétiques).")

# ================================================================ 2. MARCHÉ
titre(doc, "2. Étude de marché et analyse concurrentielle", 1)
titre(doc, "2.1 Analyse de marché", 2)
apercu(doc, "Présenter une analyse chiffrée du marché de l'aide au diagnostic génétique / e-santé : "
       "taille du marché, dynamique de croissance, cibles (CHU, hôpitaux, laboratoires), contraintes "
       "réglementaires (dispositif médical, marquage CE, règlement IA européen). Sources reconnues "
       "et datées de moins de 5 ans, citées en note de bas de page.")
titre(doc, "2.2 Analyses concurrentielles", 2)
para(doc, "Trois concurrents sont analysés selon une trame commune (positionnement, offre, forces, "
     "faiblesses), dont deux concurrents directs et un concurrent indirect.")
apercu(doc, "Concurrent direct 1, concurrent direct 2 (ex. éditeurs de logiciels d'aide au "
       "diagnostic génétique), concurrent indirect (ex. séquençage haut débit / laboratoires). "
       "Rédiger chaque analyse à la suite (pas sous forme de tableau), puis conclure sur les points "
       "positifs et négatifs retenus en comparaison avec GénoPed.")

# ================================================================ 3. PROBLÉMATIQUE
titre(doc, "3. Problématique et définition du besoin", 1)
para(doc, "Le goulot d'étranglement de GénoPed est humain : chaque dossier est analysé manuellement "
     "par un généticien senior, ce qui prend 3 à 6 jours. Avec 4 800 dossiers par an et seulement "
     "5 généticiens, l'entreprise ne peut pas absorber sa croissance et les délais s'allongent.")
para(doc, "Le besoin est donc un outil de pré-orientation automatique, capable d'analyser un dossier "
     "entrant et de suggérer instantanément les troubles génétiques les plus probables, avant même "
     "l'intervention du spécialiste. L'outil ne remplace pas le généticien : il lui prépare le "
     "terrain et lui fait gagner du temps sur les cas simples pour qu'il se concentre sur les cas "
     "complexes.")
para(doc, "Une contrainte méthodologique majeure structure le projet : les résultats de tests "
     "génétiques (noms des gènes, statut du panel) sont exclus des variables du modèle. Obtenus "
     "après le diagnostic, ils constitueraient une fuite d'information rendant le modèle "
     "artificiellement performant mais inutile en situation réelle — où l'outil doit précisément "
     "aider à décider quel test prescrire.")
para(doc, "Problématique retenue :", gras=True)
para(doc, "« Dans quelle mesure un modèle de classification supervisée, entraîné uniquement sur les "
     "données cliniques disponibles avant tout test génétique, peut-il orienter le diagnostic des "
     "troubles génétiques de l'enfant en proposant au praticien une liste courte d'hypothèses "
     "suffisamment fiable pour cibler les examens de confirmation ? »", italique=True)
para(doc, "Sur le plan fonctionnel, la solution est une application web par laquelle un praticien "
     "saisit les signes cliniques d'un enfant et reçoit les trois diagnostics les plus probables, "
     "leur justification clinique et les examens de confirmation à prescrire en priorité.")

# ================================================================ 4. GESTION DE PROJET
titre(doc, "4. Gestion de projet", 1)

titre(doc, "4.1 Pilotage du projet et planification", 2)
para(doc, "Méthode retenue : une approche agile de type Scrum adaptée à un projet data, organisée en "
     "sprints de deux semaines. Ce choix se justifie par la nature exploratoire du projet (les "
     "performances d'un modèle ne sont pas connues à l'avance) et par le besoin de points de "
     "validation réguliers avec le client (GénoPed) et le référent métier (Dr. Morin).")
para(doc, "Rétroplanning (livrables et échéances) :", gras=True)
tableau(doc, ["Phase", "Période", "Livrable"], [
    ["Cadrage & analyse du besoin", "Janvier 2026", "Cahier des charges, indicateurs de performance"],
    ["Exploration & modélisation", "Mars 2026", "Notebooks EDA + modélisation, modèle retenu"],
    ["Développement applicatif", "Mai 2026", "Application Flask + base de données"],
    ["Livraison & dossier", "Août 2026", "Code source, documentation, soutenance"],
])
para(doc, "Un diagramme de Gantt reprend ce rétroplanning et matérialise les dépendances entre "
     "tâches (l'exploration précède la modélisation, qui précède le développement).")
apercu(doc, "Insérer une capture du diagramme de Gantt (réalisable sous Excel, GanttProject ou "
       "Notion). Ajouter les tableaux de bord de suivi (avancement des sprints, burndown chart).")
para(doc, "Estimation du budget :", gras=True)
tableau(doc, ["Poste", "Détail", "Coût estimé"], [
    ["Consultant Data/IA", "4 mois × 1 ETP", "à chiffrer"],
    ["Infrastructure (serveur HDS)", "Hébergement + base PostgreSQL", "à chiffrer"],
    ["Licences / outils", "Environnement de développement", "à chiffrer"],
    ["Total prévisionnel", "", "à chiffrer"],
])

titre(doc, "4.2 Veille technologique, sectorielle et réglementaire", 2)
tableau(doc, ["Source", "Type de veille", "Outil", "Canal / fréquence", "Dernière mise à jour"], [
    ["Orphanet, HAS", "Sectorielle / médicale", "Newsletter, RSS", "Mensuelle", "Juillet 2026"],
    ["arXiv, Papers with Code", "Technologique (ML santé)", "Alertes RSS", "Hebdomadaire", "Juillet 2026"],
    ["CNIL, Règlement IA (UE)", "Juridique / réglementaire", "Site officiel, alertes", "Mensuelle", "Juillet 2026"],
    ["Concurrents MedTech", "Concurrentielle", "Google Alerts, LinkedIn", "Mensuelle", "Juillet 2026"],
])

titre(doc, "4.3 Cartographie des risques", 2)
tableau(doc, ["Risque", "Nature", "Gravité", "Mesure de maîtrise"], [
    ["Fuite d'information (gènes)", "Qualité / validité", "Élevée", "Exclusion des variables de fuite"],
    ["Déséquilibre des classes", "Qualité des données", "Élevée", "F1 macro + class_weight='balanced'"],
    ["Données de santé sensibles", "Sécurité / RGPD", "Élevée", "Pseudonymisation, hébergement HDS"],
    ["Sur-confiance dans l'IA", "Éthique / clinique", "Élevée", "Top-3, avertissements, décision au médecin"],
    ["Biais démographiques", "Éthique", "Moyenne", "Analyse des biais (EDA), variables neutres écartées"],
])
para(doc, "Enjeux environnementaux et sociétaux : le modèle retenu (Forêt aléatoire) a une empreinte "
     "de calcul modeste, entraîné en quelques secondes sur un poste standard, sans recours à des "
     "infrastructures GPU énergivores. Sur le plan sociétal, l'outil vise à réduire les inégalités "
     "d'accès au diagnostic entre établissements dotés ou non d'un généticien.")
para(doc, "Charte éthique (extraits) :", gras=True)
puce(doc, "L'outil est une aide à l'orientation : la décision diagnostique appartient au médecin.")
puce(doc, "Transparence : chaque prédiction est accompagnée de sa justification et de ses limites.")
puce(doc, "Non-discrimination : les variables sans lien clinique (sexe, origine) ne sont pas "
      "exploitées comme critères déterminants ; leur neutralité a été vérifiée.")
puce(doc, "Protection des données : aucune donnée identifiante n'est stockée ; traçabilité "
      "pseudonymisée et non altérable.")

# ================================================================ 5. EXPLOITATION DES DONNÉES
titre(doc, "5. Exploitation des données", 1)

titre(doc, "5.1 Identification des sources de données", 2)
para(doc, "La source principale est un fichier CSV (genetic_disorders_pediatric_v2.csv) extrait de "
     "la base PostgreSQL de GénoPed, l'un des trois formats admis (CSV, XLSX, JSON). Volumétrie : "
     "18 943 observations, 30 variables, soit environ 4,2 Mo après nettoyage.")
para(doc, "Typographie des données : 6 variables quantitatives (âges, numérations sanguines, nombre "
     "d'avortements), 5 variables binaires (symptômes) et 18 variables qualitatives (antécédents, "
     "facteurs génétiques), plus la variable cible qualitative à 9 modalités.")
para(doc, "Dictionnaire de données (extrait) :", gras=True)
lignes_dico = [[nom, info["type"], info["nb_modalites"], info["exemple"]]
               for nom, info in list(dico.items())]
tableau(doc, ["Variable", "Type", "Nb modalités", "Exemple"], lignes_dico)
para(doc, "Respect de la réglementation : les données sont pseudonymisées (aucun identifiant "
     "patient), hébergées sur infrastructure HDS. Points de contrôle : vérification de l'absence de "
     "données directement identifiantes, minimisation (seules les variables utiles au diagnostic "
     "sont conservées), traçabilité des accès.")

titre(doc, "5.2 Manipulation des tables et analyse", 2)
para(doc, "Installation de la base et affichage des tables : la base PostgreSQL genoped contient la "
     "table patients (18 943 lignes) et la table consultations (traçabilité des analyses). "
     "L'installation et l'affichage des tables par défaut sont documentés par captures d'écran.")
apercu(doc, "Insérer les captures pgAdmin : liste des tables, aperçu de la table patients, "
       "exécution d'une requête SELECT COUNT(*).")

para(doc, "Analyse des valeurs manquantes et des incohérences :", gras=True)
para(doc, "Les valeurs manquantes ont été traitées par imputation : la médiane pour les variables "
     "quantitatives (valeur centrale robuste aux valeurs extrêmes) et la modalité « Inconnu » pour "
     "les variables qualitatives (choix honnête, qui n'invente pas d'information). Les incohérences "
     "détectées (unité d'âge, doublons) ont été corrigées lors du nettoyage.")

para(doc, "Sécurisation des données :", gras=True)
puce(doc, "Rôle applicatif dédié (genoped_app) au lieu du super-utilisateur : principe du moindre "
      "privilège.")
puce(doc, "Lecture seule sur les dossiers sources ; l'application ne peut ni modifier ni supprimer "
      "une consultation (audit non altérable).")
puce(doc, "Révocation des droits publics par défaut ; rôle de reporting séparé en lecture seule.")
puce(doc, "Pseudonymisation systématique ; hébergement HDS.")
apercu(doc, "Insérer la capture d'écran du tableau des permissions (résultat du script "
       "sql/securisation.sql).")

para(doc, "Analyse par apprentissage supervisé :", gras=True)
para(doc, "Le jeu de données a été scindé en 75 % d'entraînement et 25 % de test (au-delà du minimum "
     "de 70 % requis), avec stratification pour préserver la répartition des classes. Cinq "
     "algorithmes de classification ont été comparés : régression logistique, arbre de décision, "
     "forêt aléatoire, K plus proches voisins (KNN) et XGBoost. Les scripts de traitement, "
     "d'entraînement et de test sont commentés dans les notebooks.")

para(doc, "Évaluation des performances :", gras=True)
para(doc, "Le déséquilibre des classes (de 25,8 % à 0,5 %) impose de privilégier le F1 macro plutôt "
     "que l'accuracy seule (un modèle prédisant toujours la classe majoritaire atteindrait déjà "
     "25,8 % d'accuracy). Résultats après optimisation des hyperparamètres :")
tableau(doc, ["Modèle", "Accuracy", "F1 macro"], [
    ["Forêt aléatoire (retenu)", "0,61", "0,42"],
    ["Régression logistique", "0,50", "0,40"],
    ["XGBoost", "0,62", "0,40"],
    ["Arbre de décision", "0,49", "0,38"],
    ["KNN", "0,54", "0,31"],
])
para(doc, "Comparaison des temps d'exécution des requêtes SQL (tables non optimisées vs optimisées "
     "par index) :")
tableau(doc, ["Requête", "Non optimisé", "Optimisé", "Gain"], [
    ["Filtre sur la maladie", "7,96 ms", "0,18 ms", "97,7 %"],
    ["Filtre maladie + symptôme", "7,59 ms", "0,13 ms", "98,3 %"],
    ["Agrégat par maladie", "9,42 ms", "1,28 ms", "86,4 %"],
])

para(doc, "Documentation technique du modèle retenu :", gras=True)
para(doc, "Le modèle final est une Forêt aléatoire (400 arbres, min_samples_leaf=10, "
     "max_features='sqrt', class_weight='balanced'). Il atteint sur le jeu de test une accuracy de "
     "57 %, un F1 macro de 0,42 et surtout une Top-3 accuracy de 87 % : la bonne maladie figure "
     "parmi les trois hypothèses proposées dans près de 9 cas sur 10, ce qui correspond à l'usage "
     "métier d'un outil d'orientation. Les techniques SMOTE et différentes pondérations de classes "
     "ont été testées et écartées (elles dégradaient le F1 macro).")

para(doc, "Prise en compte des règles d'éthique :", gras=True)
puce(doc, "Exclusion des variables de fuite pour un usage clinique honnête.")
puce(doc, "Vérification de la neutralité des variables démographiques (pas de discrimination).")
puce(doc, "Restitution en top-3 assortie d'avertissements sur les limites et les maladies rares.")

para(doc, "Tableau de suivi des problématiques techniques rencontrées :", gras=True)
tableau(doc, ["N°", "Date", "Problématique", "Date résolution", "Solution"], [
    ["1", "Mars 2026", "Fuite d'information (noms de gènes)", "Mars 2026",
     "Exclusion des 3 variables de fuite"],
    ["2", "Mars 2026", "Classes très déséquilibrées", "Mars 2026",
     "F1 macro + class_weight='balanced'"],
    ["3", "Mai 2026", "Reproductibilité des résultats", "Mai 2026",
     "Modèle final figé sur ses hyperparamètres"],
    ["4", "Mai 2026", "Incompatibilité imbalanced-learn / scikit-learn", "Mai 2026",
     "Mise à jour d'imbalanced-learn"],
    ["5", "Juin 2026", "Requête d'agrégat ralentie par l'index simple", "Juin 2026",
     "Ajout d'un index couvrant"],
])

# ================================================================ 6. APPLICATION
titre(doc, "6. Développement de l'application", 1)
para(doc, "L'application web est développée avec le framework Flask et héberge le modèle "
     "d'apprentissage supervisé. Elle est accessible localement à l'adresse http://127.0.0.1:5000 "
     "(et destinée à un déploiement sur serveur distant pour l'URL publique).")
titre(doc, "6.1 Partie front (visuelle)", 2)
para(doc, "L'interface se compose d'une page d'accueil de présentation, d'un formulaire de saisie "
     "organisé en étapes (symptômes, facteurs génétiques, informations sur l'enfant, grossesse, "
     "antécédents, examens), et d'une page de résultats. Le formulaire guide le praticien pas à pas "
     "et n'exige que les informations réellement connues.")
titre(doc, "6.2 Partie back (fonctionnalités)", 2)
para(doc, "À la soumission, l'application reconstitue le dossier, applique le pipeline de "
     "prétraitement et le modèle, puis calcule les probabilités des 9 maladies. Elle restitue les "
     "trois plus probables, chacune assortie de sa description, des signes cliniques en faveur et en "
     "défaveur (issus de la signature clinique réelle des données), et des examens de confirmation à "
     "prescrire. Chaque analyse est enregistrée dans la table consultations à des fins de "
     "traçabilité, sans donnée identifiante. Un back-office administrateur affiche l'historique et "
     "les statistiques d'usage.")
titre(doc, "6.3 Tests et déploiement", 2)
para(doc, "L'application a été testée sur serveur local (routes, prédiction de bout en bout, "
     "enregistrement en base, protection du back-office). Le déploiement sur un serveur distant "
     "fournira l'URL publique du livrable.")
apercu(doc, "Insérer les captures de l'application (accueil, formulaire, résultats, back-office) et "
       "l'URL publique une fois le déploiement réalisé.")
titre(doc, "6.4 Protection des données personnelles", 2)
para(doc, "Aucune donnée identifiante n'est collectée ni stockée ; les saisies servent uniquement au "
     "calcul. La traçabilité est pseudonymisée. En production, l'application est prévue pour un "
     "hébergement HDS avec HTTPS et durée de conservation limitée.")
titre(doc, "6.5 Accessibilité", 2)
para(doc, "L'interface respecte les bonnes pratiques d'accessibilité : structure sémantique, "
     "libellés associés à chaque champ, navigation au clavier avec focus visible, contrastes "
     "suffisants, lien d'évitement et zone de résultats annoncée aux lecteurs d'écran (aria-live).")

# ================================================================ 7. CONCLUSION
titre(doc, "7. Conclusion", 1)
para(doc, "Le projet GP AI aboutit à un système complet d'aide à l'orientation diagnostique : d'un "
     "jeu de données brut à une application web utilisable par un praticien, en passant par une "
     "analyse exploratoire rigoureuse et une modélisation comparée et optimisée. Le choix "
     "méthodologique fort — exclure les variables de fuite — garantit que l'outil reste utile en "
     "situation réelle, et l'évaluation en Top-3 (87 %) montre sa pertinence comme aide à la "
     "décision plutôt que comme substitut au médecin.")
para(doc, "Contraintes et limites : les performances restent modestes sur les maladies rares (peu "
     "d'exemples, symptômes qui se recoupent), et les données proviennent exclusivement de patients "
     "déjà atteints (absence de témoins sains), ce qui borne l'outil à l'orientation différentielle.")
para(doc, "Évolutions possibles : enrichir le jeu de données sur les pathologies rares, intégrer une "
     "classification hiérarchique (famille de maladie puis sous-type), ajouter un module de "
     "confirmation post-test, et déployer l'outil en conditions réelles avec un suivi des "
     "performances dans le temps.")

doc.save(SORTIE)
print("Rapport genere :", SORTIE)
print("Paragraphes :", len(doc.paragraphs), "| Tableaux :", len(doc.tables))
